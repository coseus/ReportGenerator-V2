
from __future__ import annotations

import base64
from collections import Counter
from io import BytesIO
from pathlib import Path
from tempfile import NamedTemporaryFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from util.charting import risk_trend_png, severity_distribution_png
from util.helpers import normalize_images
from util.i18n import t

SEV_ORDER = ["Critical", "High", "Moderate", "Low", "Informational"]
SEV_COLORS = {
    "Critical": "A61B1B",
    "High": "D35400",
    "Moderate": "B9770E",
    "Low": "1F618D",
    "Informational": "5D6D7E",
}
SECTION_DEFAULTS = {
    "section_1_0_confidentiality_and_legal": (
        "This penetration testing report contains confidential information intended solely for the client organization. "
        "Unauthorized access, distribution, disclosure, or copying of this document or any information contained herein is strictly prohibited. "
        "All findings, methodologies, and artifacts are the intellectual property of the security testing provider unless otherwise stated."
    ),
    "section_1_1_confidentiality_statement": (
        "Findings are provided for informational purposes only and represent the system state at the time of testing only. "
        "The client is solely responsible for implementing and verifying any remediation actions. "
        "The testing team disclaims all liability for any damages resulting from the use of this report or the authorized testing activities."
    ),
    "section_1_2_disclaimer": (
        "This penetration test was conducted exclusively in accordance with the Rules of Engagement and Statement of Work signed by the Client. "
        "No warranties of any kind, express or implied, are provided. The Testing Team and its personnel shall not be held liable for any direct, indirect, incidental, consequential, or punitive damages arising from the use or misuse of this report, its findings, or any actions taken as a result thereof. The report is delivered as is."
    ),
}


def _safe_text(value) -> str:
    return "" if value is None else str(value)


def _section_value(report: dict, key: str, default: str = "") -> str:
    value = report.get(key)
    if not value:
        sections = report.get("sections") or {}
        value = sections.get(key)
    if isinstance(value, dict):
        value = value.get("content") or value.get("text") or value.get("value") or ""
    return str(value or default or "").strip()


def _header_label(report: dict) -> str:
    client = _safe_text(report.get("client", "Client")).strip() or "Client"
    project = _safe_text(report.get("project", "")).strip()
    return f"{client} · {project}" if project else client


def _normalize_walkthrough_text(text: str) -> list[str]:
    text = "" if text is None else str(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" n ", "\n- ")
    text = text.replace("•", "\n- ")
    lines = []
    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line.startswith("n "):
            line = "- " + line[2:].strip()
        lines.append(line)
    return lines


def _risk_rank(value: str) -> int:
    try:
        return SEV_ORDER.index(value)
    except ValueError:
        return len(SEV_ORDER)


def _findings_sorted(report: dict) -> list[dict]:
    return sorted(report.get("findings", []) or [], key=lambda f: (_risk_rank(f.get("severity", "Informational")), (f.get("title") or "").lower()))


def _compute_summary(report: dict) -> dict:
    findings = _findings_sorted(report)
    counts = Counter()
    evidence_items = 0
    for finding in findings:
        sev = finding.get("severity", "Informational")
        if sev not in SEV_ORDER:
            sev = "Informational"
        counts[sev] += 1
        evidence_items += len(finding.get("images", []) or []) + (1 if finding.get("code") else 0)
    for item in report.get("detailed_walkthrough", []) or []:
        evidence_items += len(item.get("images", []) or []) + (1 if item.get("code") else 0)
    for item in report.get("additional_reports", []) or []:
        evidence_items += len(item.get("images", []) or []) + (1 if item.get("code") else 0)
    return {
        "findings": findings,
        "counts": {sev: counts.get(sev, 0) for sev in SEV_ORDER},
        "total": sum(counts.values()),
        "highest_severity": next((sev for sev in SEV_ORDER if counts[sev] > 0), "Informational"),
        "evidence_items": evidence_items,
    }


def _ensure_styles(doc: Document, accent: str):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)

    if "Report Title" not in styles:
        s = styles.add_style("Report Title", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Title"]
        s.font.name = "Aptos Display"
        s.font.size = Pt(24)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(accent)

    if "Report Heading 1" not in styles:
        s = styles.add_style("Report Heading 1", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Heading 1"]
        s.font.name = "Aptos Display"
        s.font.size = Pt(18)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(accent)

    if "Report Heading 2" not in styles:
        s = styles.add_style("Report Heading 2", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Heading 2"]
        s.font.name = "Aptos"
        s.font.size = Pt(12.5)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string("111827")

    if "Report Code" not in styles:
        s = styles.add_style("Report Code", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Normal"]
        s.font.name = "Consolas"
        s.font.size = Pt(8.5)

    if "Report Caption" not in styles:
        s = styles.add_style("Report Caption", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = styles["Normal"]
        s.font.name = "Aptos"
        s.font.size = Pt(9)
        s.font.italic = True
        s.font.color.rgb = RGBColor.from_string("4B5563")


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def _insert_toc(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Update field in Word to refresh the table of contents."
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)


def _add_header_footer(doc: Document, report: dict, accent: str):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run(_header_label(report))
    hr.font.size = Pt(8.5)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor.from_string(accent)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Page ")
    fr.font.size = Pt(8)
    _set_page_number(fp)


def _write_picture(doc: Document, image_b64: str, width_inches=5.9):
    try:
        raw = base64.b64decode(image_b64)
        with NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(raw)
            tmp_path = tmp.name
        doc.add_picture(tmp_path, width=Inches(width_inches))
        Path(tmp_path).unlink(missing_ok=True)
        return True
    except Exception:
        return False


def _write_picture_bytes(doc: Document, image_bytes: bytes, width_inches=5.9):
    try:
        with NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(image_bytes)
            tmp_path = tmp.name
        doc.add_picture(tmp_path, width=Inches(width_inches))
        Path(tmp_path).unlink(missing_ok=True)
        return True
    except Exception:
        return False


def _add_cover(doc: Document, report: dict, summary: dict, accent: str, report_variant: str):
    top_space = doc.add_paragraph()
    top_space.paragraph_format.space_after = Pt(8)
    if report.get("logo_b64"):
        _write_picture(doc, report["logo_b64"], width_inches=1.2)
        try:
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[-1].paragraph_format.space_after = Pt(10)
        except Exception:
            pass
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rk = kicker.add_run(t(report, f"cover_kicker_{report_variant if report_variant in {'executive', 'technical'} else 'combined'}"))
    rk.bold = True
    rk.font.color.rgb = RGBColor.from_string(accent)
    rk.font.size = Pt(11)
    kicker.paragraph_format.space_after = Pt(6)
    title = doc.add_paragraph(style="Report Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    title.add_run(t(report, "report_title"))

    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    fields = [
        (t(report, "client"), report.get("client", "N/A")),
        (t(report, "project"), report.get("project", "N/A")),
        (t(report, "assessment_date"), report.get("date", "N/A")),
        (t(report, "version"), report.get("version", "1.0")),
        (t(report, "lead_tester"), report.get("tester", "N/A")),
    ]
    for i, (label, value) in enumerate(fields):
        meta.cell(i, 0).text = label
        meta.cell(i, 1).text = _safe_text(value)
        _set_cell_shading(meta.cell(i, 0), "EEF2F7")
        for run in meta.cell(i, 0).paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    snap = doc.add_table(rows=2, cols=4)
    snap.style = "Table Grid"
    cells = [
        (0, 0, "Total Findings"), (0, 1, str(summary["total"])), (0, 2, "Highest Severity"), (0, 3, summary["highest_severity"]),
        (1, 0, "Evidence Items"), (1, 1, str(summary["evidence_items"])), (1, 2, "Confidentiality"), (1, 3, "Restricted distribution"),
    ]
    for r, c, value in cells:
        snap.cell(r, c).text = value
        if c % 2 == 0:
            _set_cell_shading(snap.cell(r, c), "F8FAFC")
            for run in snap.cell(r, c).paragraphs[0].runs:
                run.bold = True

    note = doc.add_paragraph()
    note.add_run(t(report, "confidential_note")).italic = True
    note.paragraph_format.space_before = Pt(10)
    doc.add_page_break()


def _add_toc(doc: Document):
    h = doc.add_paragraph(style="Report Heading 1")
    h.add_run(t(doc._report_ctx, "table_of_contents"))
    p = doc.add_paragraph()
    _insert_toc(p)
    doc.add_page_break()


def _add_heading(doc: Document, text: str, level: int = 1):
    style = "Report Heading 1" if level == 1 else "Report Heading 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def _add_para(doc: Document, text: str = ""):
    return doc.add_paragraph(_safe_text(text))


def _add_bullets(doc: Document, items: list[str], report: dict):
    if not items:
        doc.add_paragraph(t(report, "no_items"))
        return
    for item in items:
        doc.add_paragraph(_safe_text(item), style="List Bullet")


def _add_front_matter(doc: Document, report: dict):
    sec_10 = _section_value(report, "section_1_0_confidentiality_and_legal", SECTION_DEFAULTS["section_1_0_confidentiality_and_legal"])
    sec_11 = _section_value(report, "section_1_1_confidentiality_statement", SECTION_DEFAULTS["section_1_1_confidentiality_statement"])
    sec_12 = _section_value(report, "section_1_2_disclaimer", SECTION_DEFAULTS["section_1_2_disclaimer"])
    sec_13 = _section_value(report, "section_1_3_contact_information", "")

    _add_heading(doc, t(report, "legal"), 1)
    _add_para(doc, sec_10)
    _add_heading(doc, t(report, "conf_statement"), 2)
    _add_para(doc, sec_11)
    _add_heading(doc, t(report, "disclaimer"), 2)
    _add_para(doc, sec_12)
    _add_heading(doc, t(report, "contact_info"), 2)
    if sec_13:
        _add_para(doc, sec_13)

    contacts = report.get("contacts", []) or []
    if contacts:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Name"
        hdr[1].text = "Title"
        hdr[2].text = "Contact"
        for cell in hdr:
            _set_cell_shading(cell, "EEF2F7")
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for c in contacts:
            row = table.add_row().cells
            row[0].text = _safe_text(c.get("name", ""))
            row[1].text = _safe_text(c.get("title", ""))
            row[2].text = _safe_text(c.get("contact", ""))
    elif not sec_13:
        _add_para(doc, t(report, "no_contacts"))
    doc.add_page_break()


def _add_final_page(doc: Document, report: dict):
    doc.add_page_break()
    space = doc.add_paragraph()
    space.paragraph_format.space_after = Pt(140)
    if report.get("logo_b64"):
        _write_picture(doc, report["logo_b64"], width_inches=1.1)
        try:
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
    p = doc.add_paragraph(style="Report Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(t(report, "end_of_report"))
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(t(report, "last_page"))
    tpar = doc.add_paragraph()
    tpar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tpar.add_run(t(report, "closing_text"))


def _add_findings_summary(doc: Document, summary: dict, report: dict):
    _add_heading(doc, t(report, "findings_summary"), 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = [t(report, "severity"), t(report, "count"), t(report, "priority")]
    for i, text in enumerate(hdr):
        table.cell(0, i).text = text
        _set_cell_shading(table.cell(0, i), "EEF2F7")
    priority = {
        "Critical": t(report, "priority_critical"),
        "High": t(report, "priority_high"),
        "Moderate": t(report, "priority_moderate"),
        "Low": t(report, "priority_low"),
        "Informational": t(report, "priority_info"),
    }
    for sev in SEV_ORDER:
        row = table.add_row().cells
        row[0].text = sev
        row[1].text = str(summary["counts"][sev])
        row[2].text = priority[sev]
        _set_cell_shading(row[0], SEV_COLORS[sev])
        for run in row[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)


def _add_charts(doc: Document, report: dict, summary: dict):
    if not report.get("include_charts", True):
        return
    _add_heading(doc, t(report, "risk_charts"), 2)
    if _write_picture_bytes(doc, severity_distribution_png(summary["counts"]), width_inches=6.2):
        try:
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
        cap = doc.add_paragraph(style="Report Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(t(report, "severity_distribution"))
    if _write_picture_bytes(doc, risk_trend_png(summary["findings"]), width_inches=6.2):
        try:
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
        cap = doc.add_paragraph(style="Report Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(t(report, "risk_trend"))


def _add_overview(doc: Document, report: dict, summary: dict, report_variant: str):
    _add_heading(doc, t(report, "exec_overview") if report_variant in {"executive", "combined"} else t(report, "engagement_overview"), 1)
    _add_para(doc, report.get("executive_summary") or report.get("assessment_overview") or "No executive summary provided.")
    _add_findings_summary(doc, summary, report)
    _add_charts(doc, report, summary)
    _add_heading(doc, t(report, "assessment_details"), 2)
    for label, value in [
        ("Assessment Overview", report.get("assessment_overview", "")),
        ("Assessment Details", report.get("assessment_details", "")),
        ("Scope", report.get("scope", "")),
        ("Scope Exclusions", report.get("scope_exclusions", "")),
        ("Client Allowances", report.get("client_allowances", "")),
    ]:
        if value:
            p = doc.add_paragraph()
            r = p.add_run(label + ": ")
            r.bold = True
            p.add_run(_safe_text(value))
    attack_path = report.get("attack_path", []) or []
    if attack_path:
        p = doc.add_paragraph()
        p.add_run(t(report, "attack_path") + ": ").bold = True
        for item in attack_path:
            doc.add_paragraph(_safe_text(item), style="List Bullet")


def _finding_meta(finding: dict) -> str:
    meta = []
    for label, list_key, key in (("Hosts", "hosts", "host"), ("Ports", "ports", "port"), ("CVEs", "cves", "cve"), ("CWEs", "cwes", "cwe")):
        vals = finding.get(list_key)
        if not isinstance(vals, list):
            vals = [finding.get(key)] if finding.get(key) else []
        vals = [str(v).strip() for v in vals if str(v).strip()]
        if vals:
            meta.append(f"{label}: {', '.join(vals)}")
    for label, key in (("Protocol", "protocol"), ("CVSS", "cvss")):
        if finding.get(key):
            meta.append(f"{label}: {_safe_text(finding.get(key))}")
    return " | ".join(meta)


def _add_finding(doc: Document, idx: int, finding: dict, prefix: str, report: dict):
    title = finding.get("title") or f"Finding {idx}"
    _add_heading(doc, f"{prefix}.{idx} {title}", 2)

    p2 = doc.add_paragraph()
    p2.add_run(t(report, "severity") + ": ").bold = True
    p2.add_run(_safe_text(finding.get("severity", "Informational")))
    p2.paragraph_format.space_after = Pt(6)

    meta_line = _finding_meta(finding)
    if meta_line:
        mp = doc.add_paragraph(meta_line)
        mp.paragraph_format.space_after = Pt(6)

    for label_key, key in [
        ("description", "description"),
        ("likelihood", "likelihood"),
        ("impact", "impact"),
        ("tools_used", "tools_used"),
        ("recommendation", "recommendation"),
        ("references", "references"),
    ]:
        value = finding.get(key)
        if value:
            p = doc.add_paragraph()
            p.add_run(t(report, label_key) + ": ").bold = True
            p.add_run(_safe_text(value))
            p.paragraph_format.space_after = Pt(4)

    if finding.get("code"):
        label_p = doc.add_paragraph()
        label_p.add_run(t(report, "evidence_output")).bold = True
        label_p.paragraph_format.space_after = Pt(6)
        code_p = doc.add_paragraph(_safe_text(finding.get("code")), style="Report Code")
        code_p.paragraph_format.space_after = Pt(14)

    for image_item in normalize_images(finding.get("images"), default_prefix=title):
        if _write_picture(doc, image_item["data"]):
            try:
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.paragraphs[-1].paragraph_format.space_after = Pt(0)
            except Exception:
                pass
            if image_item.get("name"):
                cap = doc.add_paragraph(style="Report Caption")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_before = Pt(7)
                cap.paragraph_format.space_after = Pt(12)
                cap.add_run(_safe_text(image_item["name"]))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)


def _add_remediation(doc: Document, report: dict, section_number: str):
    _add_heading(doc, t(report, "remediation", n=section_number), 1)
    mapping = [
        (t(report, "short_term", n=section_number), report.get("remediation_short", [])),
        (t(report, "medium_term", n=section_number), report.get("remediation_medium", [])),
        (t(report, "long_term", n=section_number), report.get("remediation_long", [])),
    ]
    for heading, items in mapping:
        _add_heading(doc, heading, 2)
        _add_bullets(doc, items, report)


def _add_walkthrough(doc: Document, report: dict, section_number: str):
    steps = report.get("detailed_walkthrough", []) or []
    if not steps:
        return
    _add_heading(doc, t(report, "walkthrough", n=section_number), 1)
    for idx, step in enumerate(steps, start=1):
        title = step.get("name") or step.get("title") or f"Step {idx}"
        _add_heading(doc, f"{section_number}.{idx} {title}", 2)
        for line in _normalize_walkthrough_text(step.get("description", "")):
            if line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                p = _add_para(doc, line)
                p.paragraph_format.space_after = Pt(4)
        if step.get("code"):
            label_p = doc.add_paragraph()
            label_p.add_run(t(report, "command_output")).bold = True
            label_p.paragraph_format.space_after = Pt(6)
            code_p = doc.add_paragraph(_safe_text(step.get("code")), style="Report Code")
            code_p.paragraph_format.space_after = Pt(14)
        for image_item in normalize_images(step.get("images"), default_prefix=title):
            if _write_picture(doc, image_item["data"]):
                try:
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.paragraphs[-1].paragraph_format.space_after = Pt(0)
                except Exception:
                    pass
                if image_item.get("name"):
                    cap = doc.add_paragraph(style="Report Caption")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.paragraph_format.space_before = Pt(7)
                    cap.paragraph_format.space_after = Pt(12)
                    cap.add_run(_safe_text(image_item["name"]))
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)


def _add_additional_reports(doc: Document, report: dict, section_number: str):
    extras = report.get("additional_reports", []) or []
    if not extras:
        return
    _add_heading(doc, t(report, "additional_reports", n=section_number), 1)
    for idx, extra in enumerate(extras, start=1):
        title = extra.get("name") or extra.get("title") or f"Additional Report {idx}"
        _add_heading(doc, f"{section_number}.{idx} {title}", 2)
        if extra.get("description"):
            p = _add_para(doc, extra.get("description"))
            p.paragraph_format.space_after = Pt(4)
        if extra.get("code"):
            label_p = doc.add_paragraph()
            label_p.add_run(t(report, "output")).bold = True
            label_p.paragraph_format.space_after = Pt(6)
            code_p = doc.add_paragraph(_safe_text(extra.get("code")), style="Report Code")
            code_p.paragraph_format.space_after = Pt(14)
        for image_item in normalize_images(extra.get("images"), default_prefix=title):
            if _write_picture(doc, image_item["data"]):
                try:
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.paragraphs[-1].paragraph_format.space_after = Pt(0)
                except Exception:
                    pass
                if image_item.get("name"):
                    cap = doc.add_paragraph(style="Report Caption")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.paragraph_format.space_before = Pt(7)
                    cap.paragraph_format.space_after = Pt(12)
                    cap.add_run(_safe_text(image_item["name"]))
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)


def generate_docx_bytes(report: dict, report_variant: str = "technical") -> bytes:
    report_variant = (report_variant or "technical").lower()
    if report_variant not in {"technical", "executive", "combined"}:
        report_variant = "technical"

    doc = Document()
    doc._report_ctx = report
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    accent = (report.get("theme_hex") or "#ED863D").replace("#", "").upper()
    _ensure_styles(doc, accent)
    _add_header_footer(doc, report, accent)
    summary = _compute_summary(report)

    _add_cover(doc, report, summary, accent, report_variant)
    _add_front_matter(doc, report)
    _add_toc(doc)
    _add_overview(doc, report, summary, report_variant)

    if report_variant == "executive":
        _add_heading(doc, t(report, "key_findings"), 1)
        for idx, finding in enumerate(summary["findings"][: min(6, len(summary["findings"]))], start=1):
            _add_finding(doc, idx, finding, "4", report)
        _add_remediation(doc, report, "5")
    elif report_variant == "combined":
        _add_heading(doc, t(report, "key_findings"), 1)
        for idx, finding in enumerate(summary["findings"][: min(6, len(summary["findings"]))], start=1):
            _add_finding(doc, idx, finding, "4", report)
        _add_remediation(doc, report, "5")
        _add_heading(doc, t(report, "technical_findings", n="6"), 1)
        for idx, finding in enumerate(summary["findings"], start=1):
            _add_finding(doc, idx, finding, "6", report)
        _add_walkthrough(doc, report, "7")
        _add_additional_reports(doc, report, "8")
    else:
        _add_heading(doc, t(report, "technical_findings", n="4"), 1)
        for idx, finding in enumerate(summary["findings"], start=1):
            _add_finding(doc, idx, finding, "4", report)
        _add_remediation(doc, report, "5")
        _add_walkthrough(doc, report, "6")
        _add_additional_reports(doc, report, "7")

    _add_final_page(doc, report)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
